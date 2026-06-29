import json
import re
import os
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

CLAUDE_MODEL = "claude-sonnet-4-6"

claude_client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])


def extract_json_from_response(content, log_func):
    if "```json" in content:
        json_match = re.search(r'```json\s*(\{.*\})\s*```', content, re.DOTALL)
        if json_match:
            content = json_match.group(1)
    elif "```" in content:
        json_match = re.search(r'```\s*(\{.*\})\s*```', content, re.DOTALL)
        if json_match:
            content = json_match.group(1)

    if not content.strip().startswith('{'):
        json_match = re.search(r'(\{.*\})', content, re.DOTALL)
        if json_match:
            content = json_match.group(1)

    try:
        return json.loads(content)
    except json.JSONDecodeError as e:
        log_func(f"JSON Parse Error: {str(e)}")

    content_fixed = re.sub(r'[\x00-\x1f]', ' ', content)
    content_fixed = re.sub(r',\s*}', '}', content_fixed)
    content_fixed = re.sub(r',\s*]', ']', content_fixed)
    content_fixed = re.sub(r' {2,}', ' ', content_fixed)

    try:
        return json.loads(content_fixed)
    except json.JSONDecodeError:
        pass

    log_func("Attempting repair...")
    try:
        repair_content = content_fixed
        if len(repair_content) > 80000:
            repair_content = content_fixed[:80000] + "\n...[truncated]...\n" + content_fixed[-5000:]

        repair_response = claude_client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=16000,
            temperature=0,
            messages=[
                {"role": "user", "content": f"The following JSON is malformed. Fix it so it parses correctly. Return ONLY the fixed JSON, nothing else. No code fences, no explanation. Common issues: unescaped quotes inside strings, missing commas, trailing commas, control characters.\n\n{repair_content}"},
                {"role": "assistant", "content": "{"},
            ]
        )

        repaired = "{" + repair_response.content[0].text
        last_brace = repaired.rfind('}')
        if last_brace > 0:
            repaired = repaired[:last_brace + 1]

        result = json.loads(repaired)
        log_func("JSON repair successful!")
        return result
    except Exception as repair_err:
        log_func(f"JSON repair failed: {str(repair_err)}")

    log_func(f"ERROR: Could not parse JSON after all fixes")
    log_func(f"Full response length: {len(content)}")
    return None


def build_ai_prompt(deponent, transcript, config):
    return f"""Analyze this deposition transcript of {deponent}. Extract COMPREHENSIVE information using EXACT phrases from transcript.

CRITICAL RULES:
1. NO apostrophes or quotes in output
2. Use EXACT facts: dates, names, numbers from testimony
3. DO NOT add deponent name to statements
4. DO NOT paraphrase - copy exact phrases
5. Include ALL weaknesses, limitations, contradictions

Create JSON response:

{{
  "abstract": "Write 4-5 detailed paragraphs (200+ words). Para 1: Full background, credentials, education, work history with exact dates and titles. Para 2: This specific case - parties names, loss date, inspection date, what was inspected, timeline details. Para 3: Detailed methodology - what they did, what they did NOT do, equipment used, time spent, measurements taken or not taken. Para 4: Key findings with exact observations and claims made. Para 5: Limitations acknowledged, payment terms, prior testimony history, any disqualifications or weaknesses revealed.",

  "key_admissions": [
    {{
      "heading": "Brief topic (3-6 words)",
      "summary": "Complete sentence with EXACT facts: specific dates, names, numbers. Focus on damaging admissions, limitations, contradictions, lack of documentation, contingency fees, missing evidence, training gaps, disqualifications.",
      "cite_start": "page:line",
      "cite_end": "page:line"
    }}
  ],

  "sections": [
    {{
      "title": "Detailed section heading describing content (8-20 words)",
      "range_start": "page:line",
      "range_end": "page:line",
      "topics": [
        {{
          "topic": "Specific topic (3-10 words)",
          "summary": "2-3 sentences with EXACT facts from testimony. Include specific measurements, times, procedures, admissions of what was NOT done, limitations, weaknesses.",
          "cite_start": "page:line",
          "cite_end": "page:line"
        }}
      ]
    }}
  ]
}}

CRITICAL FOCUS AREAS:
- Lack of measurements, testing, documentation
- Contingency fee arrangements
- Prior testimony/disqualifications
- Training gaps/no recent training
- Methodology weaknesses
- What was NOT inspected/photographed
- Contradictions between claims and evidence
- Relying on visual inspection only
- Time spent vs scope of claims

Create {config['admissions']} admissions, {config['sections']} sections with {config['topics_per']} topics each.

TRANSCRIPT:
{transcript}"""


def build_critiquer_prompt(deponent, transcript):
    return f"""You are a senior trial attorney with 25+ years of deposition experience reviewing the transcript of {deponent}. Perform a comprehensive deposition analysis covering questioning strategy, witness behavior, contradictions, and trial preparation opportunities.

CRITICAL RULES:
1. NO apostrophes or single quotes inside string values - use the word or rephrase instead
2. Use EXACT text from the transcript when referencing questions and answers
3. Be specific and actionable - every finding must help an attorney take concrete action
4. All citations must use page:line format matching the transcript
5. Rate every finding as High, Medium, or Low priority

Create JSON response with ALL of these sections:

{{
  "executive_summary": {{
    "critical_findings": [
      {{
        "finding": "Finding 1 - the single most important issue (1 sentence)",
        "cite": "page:line"
      }},
      {{
        "finding": "Finding 2 - second most important (1 sentence)",
        "cite": "page:line"
      }},
      {{
        "finding": "Finding 3 - third most important (1 sentence)",
        "cite": "page:line"
      }},
      {{
        "finding": "Finding 4 - fourth most important (1 sentence)",
        "cite": "page:line"
      }},
      {{
        "finding": "Finding 5 - fifth most important (1 sentence)",
        "cite": "page:line"
      }}
    ],
    "deponent_credibility_notes": [
      {{
        "observation": "Specific credibility observation with exact details from testimony",
        "cite": "page:line",
        "pattern": "Evasive / Cooperative / Inconsistent / Coached / Defensive / Vague"
      }}
    ]
  }},

  "contradictions": [
    {{
      "priority": "High/Medium/Low",
      "citation_1": "page:line of first statement",
      "testimony_1": "Exact text of first statement",
      "citation_2": "page:line of contradicting statement",
      "testimony_2": "Exact text of contradicting statement",
      "analysis": "Explain specifically how these statements contradict each other and why it matters for the case",
      "trial_use": "How this contradiction could be used during cross-examination at trial"
    }}
  ],

  "evasion_patterns": [
    {{
      "priority": "High/Medium/Low",
      "citation": "page:line",
      "question": "The exact question asked",
      "answer": "The exact evasive answer given",
      "evasion_type": "Non-answer / Redirect / Vague / Qualified / Memory claim / Coached response",
      "what_was_avoided": "What the witness was trying not to answer and why it likely matters",
      "pin_down_questions": "Write 2-3 follow-up questions that would force a direct answer"
    }}
  ],

  "commitment_map": [
    {{
      "priority": "High/Medium/Low",
      "citation": "page:line",
      "testimony": "The exact testimony where the commitment was made",
      "commitment": "What specific position the deponent locked themselves into",
      "locked_in": true,
      "vulnerability": "How this commitment could be used against them - what evidence or testimony could contradict this position"
    }}
  ],

  "weak_questions": [
    {{
      "priority": "High/Medium/Low",
      "citation": "page:line",
      "original_question": "The exact question text from transcript",
      "original_answer": "The exact answer text from transcript",
      "issue": "Explain why this question was weak - leading, vague, compound, allowed escape, failed to pin down, etc.",
      "suggested_improvement": "Write a better version of the question"
    }}
  ],

  "missed_follow_ups": [
    {{
      "priority": "High/Medium/Low",
      "citation": "page:line",
      "testimony": "The exact Q and A exchange where the opening occurred",
      "missed_opportunity": "What the witness revealed or implied that the attorney failed to pursue",
      "suggested_follow_up": "Write the follow-up question(s) that should have been asked"
    }}
  ],

  "topic_gaps": [
    {{
      "priority": "High/Medium/Low",
      "topic": "The topic area that was never explored",
      "why_it_matters": "Why this topic is important based on what was discussed in the deposition",
      "suggested_questions": [
        "Question 1 to explore this topic",
        "Question 2 to explore this topic",
        "Question 3 to explore this topic"
      ]
    }}
  ],

  "cross_examination_ammunition": [
    {{
      "priority": "High/Medium/Low",
      "citation": "page:line",
      "testimony": "The exact testimony that can be used",
      "how_to_use": "Specifically how to use this at trial during cross-examination",
      "setup_questions": "Questions to ask at trial before springing this testimony on the witness"
    }}
  ]
}}

ANALYSIS FOCUS:
- Contradictions: scan the ENTIRE transcript for statements that conflict with each other, even if far apart
- Evasion: flag every non-answer, redirect, memory claim, and vague response where the attorney accepted it and moved on
- Commitments: identify every firm position the deponent took that locks them in
- Weak questions: compound questions, leading at wrong times, questions that let the witness escape
- Missed follow-ups: volunteered information not pursued, partial admissions not expanded, doors opened but not walked through
- Topic gaps: areas that should have been covered based on the testimony but were never asked about
- Cross-exam ammo: testimony that can be weaponized at trial

TARGET COUNTS:
- 5 critical findings in executive summary
- 5-10 contradictions (if present - do not fabricate)
- 8-12 evasion patterns
- 8-12 commitment map entries
- 8-12 weak questions
- 8-12 missed follow-ups
- 5-8 topic gaps
- 8-12 cross-examination ammunition items

TRANSCRIPT:
{transcript}"""


def generate_critique_claude(transcript, deponent, log_func, progress_func=None):
    max_chars = 180000
    if len(transcript) > max_chars:
        transcript = transcript[:max_chars] + "\n[truncated...]"

    prompt = build_critiquer_prompt(deponent, transcript)
    expected_chars = 25000

    for attempt in range(2):
        try:
            if attempt == 0:
                log_func("Running Depo Hindsight analysis...")
            else:
                log_func("Retrying Depo Hindsight...")

            messages = [{"role": "user", "content": prompt}]

            content = ""
            with claude_client.messages.stream(
                model=CLAUDE_MODEL,
                max_tokens=32000,
                temperature=0,
                messages=messages
            ) as stream:
                for text in stream.text_stream:
                    content += text
                    if progress_func:
                        pct = min(int((len(content) / expected_chars) * 90), 90)
                        progress_func(pct, f"Analyzing... {pct}%")

            if progress_func:
                progress_func(92, "Processing response...")
            log_func(f"Hindsight response received, length: {len(content):,}")

            if progress_func:
                progress_func(95, "Parsing data...")
            result = extract_json_from_response(content, log_func)
            if result:
                if progress_func:
                    progress_func(98, "Building document...")
                return result

            log_func(f"Hindsight attempt {attempt + 1} failed to parse, trying again...")
        except Exception as e:
            log_func(f"ERROR: {str(e)}")

    return None


def generate_summary_claude(transcript, deponent, qa_pairs, log_func, config, progress_func=None):
    max_chars = 180000
    if len(transcript) > max_chars:
        transcript = transcript[:max_chars] + "\n[truncated...]"

    prompt = build_ai_prompt(deponent, transcript, config)
    expected_chars = 12000

    for attempt in range(2):
        try:
            if attempt == 0:
                log_func("Running summarizer...")
            else:
                log_func("Retrying summarizer...")

            messages = [{"role": "user", "content": prompt}]

            content = ""
            with claude_client.messages.stream(
                model=CLAUDE_MODEL,
                max_tokens=16000,
                temperature=0,
                messages=messages
            ) as stream:
                for text in stream.text_stream:
                    content += text
                    if progress_func:
                        pct = min(int((len(content) / expected_chars) * 90), 90)
                        progress_func(pct, f"Summarizing... {pct}%")

            if progress_func:
                progress_func(92, "Processing response...")
            log_func(f"Response received, length: {len(content):,}")

            if progress_func:
                progress_func(95, "Parsing data...")
            result = extract_json_from_response(content, log_func)
            if result:
                if progress_func:
                    progress_func(98, "Building document...")
                return result

            log_func(f"Attempt {attempt + 1} failed to parse, trying again...")
        except Exception as e:
            log_func(f"ERROR: {str(e)}")

    return None


def parse_transcript(text):
    qa_pairs = []
    lines = text.split('\n')
    current_page = 1
    current_speaker = None
    current_text = []
    start_line = 1
    start_page = 1

    for idx, line in enumerate(lines):
        stripped = line.strip()

        if stripped.isdigit():
            current_page = int(stripped)
            continue

        line_match = re.match(r'^\s*(\d{1,2})\s+([QA])\.?\s+(.+)$', line)

        if line_match:
            new_line = int(line_match.group(1))
            speaker = line_match.group(2)
            content = line_match.group(3).strip()

            if current_speaker and current_text:
                qa_pairs.append({
                    'page': start_page,
                    'line': start_line,
                    'speaker': current_speaker,
                    'text': ' '.join(current_text).strip()
                })

            current_speaker = speaker
            start_line = new_line
            start_page = current_page
            current_text = [content]

        elif current_speaker and stripped:
            cleaned = re.sub(r'^\d{1,2}\s+', '', stripped)

            new_qa_match = re.match(r'^([QA])\.?\s+(.+)$', cleaned)
            if new_qa_match:
                if current_text:
                    qa_pairs.append({
                        'page': start_page,
                        'line': start_line,
                        'speaker': current_speaker,
                        'text': ' '.join(current_text).strip()
                    })

                current_speaker = new_qa_match.group(1)
                start_line = start_line + 1
                start_page = current_page
                current_text = [new_qa_match.group(2).strip()]
            elif cleaned:
                current_text.append(cleaned)

    if current_speaker and current_text:
        qa_pairs.append({
            'page': start_page,
            'line': start_line,
            'speaker': current_speaker,
            'text': ' '.join(current_text).strip()
        })

    return qa_pairs


def extract_name_from_transcript(text):
    lines = text.split('\n')[:100]
    for line in lines:
        if 'deposition of' in line.lower():
            match = re.search(r'deposition of ([A-Z][a-z]+(?: [A-Z][a-z]+)+)', line, re.I)
            if match:
                return match.group(1).title()
        if re.match(r'^[A-Z][a-z]+ [A-Z][a-z]+', line):
            return line.strip()
    return "Witness"


def extract_case_caption(text):
    lines = text.split('\n')

    page_1_end = len(lines)
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped == '2' or stripped == 'Page 2':
            page_1_end = i
            break

    page_1_lines = lines[:page_1_end]

    cleaned = []
    for line in page_1_lines:
        c = re.sub(r'^\s*\d{1,2}\s+', '', line).strip()
        cleaned.append(c)

    v_index = None
    for i, line in enumerate(cleaned):
        if re.match(r'^v\.?\s*$', line, re.IGNORECASE) or re.match(r'^vs\.?\s*$', line, re.IGNORECASE):
            v_index = i
            break

    if v_index is None:
        return None

    party_labels = r'(?:plaintiff|defendant|petitioner|respondent|appellant|appellee|claimant|complainant|grievant|applicant|movant|relator|debtor|creditor|employee|employer|intervenor|cross-plaintiff|cross-defendant|counter-plaintiff|counter-defendant|third-party plaintiff|third-party defendant)'

    party1 = None
    for i in range(v_index - 1, -1, -1):
        line = cleaned[i]
        if not line:
            continue
        if re.match(rf'^{party_labels}s?[,.:]*\s*$', line, re.IGNORECASE):
            continue
        if len(line) > 3:
            party1 = line.strip(' ,.')
            break

    party2 = None
    for i in range(v_index + 1, len(cleaned)):
        line = cleaned[i]
        if not line:
            continue
        if re.match(rf'^{party_labels}s?[,.:]*\s*$', line, re.IGNORECASE):
            continue
        if len(line) > 3:
            party2 = line.strip(' ,.')
            break

    if party1 and party2:
        party1 = re.sub(rf',?\s*{party_labels}s?[,.:]*\s*$', '', party1, flags=re.IGNORECASE).strip(' ,.')
        party2 = re.sub(rf',?\s*{party_labels}s?[,.:]*\s*$', '', party2, flags=re.IGNORECASE).strip(' ,.')

        caption = f"{party1} vs. {party2}"

        if 5 < len(caption) < 300:
            return caption

    return None


def add_bookmark(paragraph, name):
    if not paragraph.runs:
        paragraph.add_run("")
    run = paragraph.runs[0]
    bookmark_id = str(abs(hash(name)) % (2**15))
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bookmark_id)
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bookmark_id)
    run._r.addprevious(start)
    run._r.addnext(end)


def add_internal_hyperlink(paragraph, display_text, anchor_name):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor_name)
    hyperlink.set(qn('w:history'), "1")
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), "Hyperlink")
    rPr.append(rStyle)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0000FF")
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), "single")
    rPr.append(color)
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = display_text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_header_footer(doc, logo_path=None):
    section = doc.sections[0]
    header = section.header
    footer = section.footer

    if logo_path and os.path.exists(logo_path):
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(1.5))

    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    footer_run = footer_para.add_run()
    footer_run._r.append(fldChar1)
    footer_run._r.append(instrText)
    footer_run._r.append(fldChar2)


def create_document(plan, deponent, citation_label, qa_pairs, case_caption=None, logo_path=None, settings=None, page_margins=None):
    doc = Document()
    doc.core_properties.author = "Rick"
    doc.core_properties.comments = "Ucrinc.com"
    doc.core_properties.title = f"{deponent} Deposition Summary"

    if settings is None:
        settings = {
            'title': {'font_size': 16, 'alignment': 'center'},
            'abstract': {'font_size': 11, 'alignment': 'justify'},
            'key_admissions': {'font_size': 12, 'alignment': 'left'},
            'toc': {'font_size': 11, 'alignment': 'left'},
            'examination': {'font_size': 11, 'alignment': 'left'},
            'transcript': {'font_size': 10, 'alignment': 'left'}
        }

    if page_margins is None:
        page_margins = {'top': 0.5, 'bottom': 0.5, 'left': 0.5, 'right': 0.5}

    def get_alignment(align_str):
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        return align_map.get(align_str.lower(), WD_ALIGN_PARAGRAPH.LEFT)

    section = doc.sections[0]
    section.top_margin = Inches(page_margins['top'])
    section.bottom_margin = Inches(page_margins['bottom'])
    section.left_margin = Inches(page_margins['left'])
    section.right_margin = Inches(page_margins['right'])

    p = doc.add_paragraph()
    p.alignment = get_alignment(settings['title']['alignment'])

    run1 = p.add_run(deponent)
    run1.bold = True
    run1.font.size = Pt(settings['title']['font_size'] + 2)

    run2 = p.add_run("\nDeposition Summary")
    run2.bold = False
    run2.font.size = Pt(settings['title']['font_size'] - 2)
    run2.font.color.rgb = RGBColor(100, 100, 100)

    if case_caption:
        p = doc.add_paragraph()
        p.alignment = get_alignment(settings['title']['alignment'])
        run = p.add_run(case_caption)
        run.italic = True
        run.font.size = Pt(11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DEPOSITION ABSTRACT")
    run.bold = True
    run.underline = True
    run.font.size = Pt(settings['abstract']['font_size'] + 3)
    doc.add_paragraph()

    abstract_text = plan.get("abstract", "No abstract available.")

    paragraphs = []
    sentences = re.split(r'(?<=[.!?])\s+', abstract_text)

    current_para = []
    sentence_count = 0

    for sentence in sentences:
        current_para.append(sentence)
        sentence_count += 1

        if sentence_count >= 4:
            paragraphs.append(' '.join(current_para))
            current_para = []
            sentence_count = 0

    if current_para:
        paragraphs.append(' '.join(current_para))

    if len(paragraphs) < 2:
        paragraphs = [abstract_text]

    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)
        p.alignment = get_alignment(settings['abstract']['alignment'])
        for run in p.runs:
            run.font.size = Pt(settings['abstract']['font_size'])
        p.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    doc.add_heading("KEY ADMISSIONS", level=1)
    admissions = plan.get("key_admissions", [])
    if admissions:
        table = doc.add_table(rows=len(admissions) + 1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].width = Inches(1.2)
        hdr_cells[1].width = Inches(2.0)
        hdr_cells[2].width = Inches(3.3)

        h1 = hdr_cells[0].paragraphs[0]
        r1 = h1.add_run("Citation")
        r1.bold = True
        r1.underline = True
        r1.font.size = Pt(settings['key_admissions']['font_size'])

        h2 = hdr_cells[1].paragraphs[0]
        r2 = h2.add_run("Heading")
        r2.bold = True
        r2.underline = True
        r2.font.size = Pt(settings['key_admissions']['font_size'])

        h3 = hdr_cells[2].paragraphs[0]
        r3 = h3.add_run("Summary")
        r3.bold = True
        r3.underline = True
        r3.font.size = Pt(settings['key_admissions']['font_size'])

        for idx, adm in enumerate(admissions):
            row = table.rows[idx + 1]
            row.cells[0].width = Inches(1.2)
            row.cells[1].width = Inches(2.0)
            row.cells[2].width = Inches(3.3)

            cite_start = adm.get("cite_start", "")
            cite_end = adm.get("cite_end", "")
            citation = f"{cite_start} - {cite_end}"
            cite_cell = row.cells[0]
            cite_p = cite_cell.paragraphs[0]

            return_bookmark = f"admission_{idx}"
            add_bookmark(cite_p, return_bookmark)

            if cite_start and cite_end:
                cite_bookmark = f"cite_{cite_start.replace(':', '_')}"
                add_internal_hyperlink(cite_p, citation, cite_bookmark)
            else:
                run = cite_p.add_run(citation)
                run.font.size = Pt(settings['key_admissions']['font_size'])

            heading_cell = row.cells[1]
            heading_p = heading_cell.paragraphs[0]
            run = heading_p.add_run(adm.get("heading", ""))
            run.bold = True
            run.font.size = Pt(settings['key_admissions']['font_size'])

            summary_cell = row.cells[2]
            summary_p = summary_cell.paragraphs[0]
            run = summary_p.add_run(adm.get("summary", ""))
            run.font.size = Pt(settings['key_admissions']['font_size'])

    doc.add_page_break()

    doc.add_heading("TABLE OF CONTENTS", level=1)
    sections = plan.get("sections", [])
    if sections:
        table = doc.add_table(rows=len(sections) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].width = Inches(1.5)
        hdr_cells[1].width = Inches(5.0)

        h1 = hdr_cells[0].paragraphs[0]
        r1 = h1.add_run("Pages")
        r1.bold = True
        r1.underline = True
        r1.font.size = Pt(settings['toc']['font_size'])

        h2 = hdr_cells[1].paragraphs[0]
        r2 = h2.add_run("Section")
        r2.bold = True
        r2.underline = True
        r2.font.size = Pt(settings['toc']['font_size'])

        for idx, section in enumerate(sections):
            row = table.rows[idx + 1]
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(5.0)
            range_start = section.get("range_start", "")
            range_end = section.get("range_end", "")
            cell_text = row.cells[0].paragraphs[0]
            run = cell_text.add_run(f"{range_start} - {range_end}")
            run.font.size = Pt(settings['toc']['font_size'])

            cell_p = row.cells[1].paragraphs[0]
            cell_p.clear()
            bookmark = f"section_{idx}"
            add_internal_hyperlink(cell_p, section.get("title", ""), bookmark)

    doc.add_page_break()

    doc.add_heading("EXAMINATION", level=1)
    for idx, section in enumerate(sections):
        doc.add_paragraph()
        p = doc.add_heading(section.get("title", ""), level=2)
        add_bookmark(p, f"section_{idx}")
        for run in p.runs:
            run.font.size = Pt(settings['examination']['font_size'] + 2)
        doc.add_paragraph()

        topics = section.get('topics', [])
        if topics:
            table = doc.add_table(rows=len(topics) + 1, cols=2)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].width = Inches(2.0)
            hdr_cells[1].width = Inches(4.5)

            h1 = hdr_cells[0].paragraphs[0]
            r1 = h1.add_run("Topic")
            r1.bold = True
            r1.underline = True
            r1.font.size = Pt(settings['examination']['font_size'])

            h2 = hdr_cells[1].paragraphs[0]
            r2 = h2.add_run("Summary")
            r2.bold = True
            r2.underline = True
            r2.font.size = Pt(settings['examination']['font_size'])

            for t_idx, topic in enumerate(topics):
                row = table.rows[t_idx + 1]
                row.cells[0].width = Inches(2.0)
                row.cells[1].width = Inches(4.5)

                cite_start = topic.get("cite_start", "")
                cite_end = topic.get("cite_end", "")
                bookmark = f"summary_{cite_start.replace(':', '_')}"
                topic_p = row.cells[0].paragraphs[0]
                add_bookmark(topic_p, bookmark)
                run = topic_p.add_run(topic.get("topic", ""))
                run.font.size = Pt(settings['examination']['font_size'])

                summary_p = row.cells[1].paragraphs[0]
                run = summary_p.add_run(topic.get("summary", ""))
                run.font.size = Pt(settings['examination']['font_size'])
                summary_p.add_run("\n")
                citation = f"{cite_start} - {cite_end}"
                cite_bookmark_link = f"cite_{cite_start.replace(':', '_')}"
                add_internal_hyperlink(summary_p, citation, cite_bookmark_link)

    doc.add_page_break()
    doc.add_heading("Transcript", level=1)

    def sort_key(qa):
        return (qa['page'], qa['line'])

    sorted_qa_pairs = sorted(qa_pairs, key=sort_key)

    admission_map = {}
    for idx, adm in enumerate(plan.get('key_admissions', [])):
        cite_start = adm.get('cite_start', '')
        cite_end = adm.get('cite_end', '')
        if cite_start and cite_end:
            key = f"{cite_start}_{cite_end}"
            admission_map[key] = idx

    table = doc.add_table(rows=len(sorted_qa_pairs) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Inches(1.2)
    hdr_cells[1].width = Inches(5.3)

    h1 = hdr_cells[0].paragraphs[0]
    run = h1.add_run("Citation")
    run.font.size = Pt(settings['transcript']['font_size'])

    h2 = hdr_cells[1].paragraphs[0]
    run = h2.add_run("Question / Answer")
    run.font.size = Pt(settings['transcript']['font_size'])

    for idx, qa in enumerate(sorted_qa_pairs):
        row = table.rows[idx + 1]
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(5.3)

        citation = f"{qa['page']}:{qa['line']}"
        cite_cell = row.cells[0]
        cite_p = cite_cell.paragraphs[0]

        bookmark_name = f"cite_{citation.replace(':', '_')}"
        add_bookmark(cite_p, bookmark_name)

        cite_run = cite_p.add_run(citation)
        cite_run.bold = True
        cite_run.font.size = Pt(settings['transcript']['font_size'])

        for key, adm_idx in admission_map.items():
            start_cite, end_cite = key.split('_', 1)
            start_parts = start_cite.split(':')
            end_parts = end_cite.split(':')
            if len(start_parts) == 2 and len(end_parts) == 2:
                try:
                    start_page, start_line = int(start_parts[0]), int(start_parts[1])
                    end_page, end_line = int(end_parts[0]), int(end_parts[1])
                    qa_page, qa_line = qa['page'], qa['line']

                    if ((qa_page > start_page or (qa_page == start_page and qa_line >= start_line)) and
                        (qa_page < end_page or (qa_page == end_page and qa_line <= end_line))):
                        cite_p.add_run(" ")
                        add_internal_hyperlink(cite_p, "[^]", f"admission_{adm_idx}")
                        break
                except (ValueError, TypeError):
                    continue

        qa_cell = row.cells[1]
        qa_p = qa_cell.paragraphs[0]
        run = qa_p.add_run(f"{qa['speaker']}    {qa['text']}")
        run.font.size = Pt(settings['transcript']['font_size'])

    for section in doc.sections:
        section.top_margin = Inches(page_margins['top'])
        section.bottom_margin = Inches(page_margins['bottom'])
        section.left_margin = Inches(page_margins['left'])
        section.right_margin = Inches(page_margins['right'])

    add_header_footer(doc, logo_path)

    return doc


def _priority_color(priority):
    p = str(priority).strip().lower()
    if p == "high":
        return RGBColor(180, 0, 0)
    elif p == "medium":
        return RGBColor(200, 130, 0)
    return RGBColor(80, 80, 80)


def _add_priority_cell(cell, priority, width):
    cell.width = width
    p = cell.paragraphs[0]
    run = p.add_run(str(priority))
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = _priority_color(priority)


def _add_critique_table_headers(table, headers, widths):
    for i, hdr_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        p = cell.paragraphs[0]
        run = p.add_run(hdr_text)
        run.bold = True
        run.underline = True
        run.font.size = Pt(10)


def create_critique_document(critique, deponent, case_caption=None, logo_path=None, page_margins=None):
    doc = Document()
    doc.core_properties.author = "Rick"
    doc.core_properties.comments = "Ucrinc.com"
    doc.core_properties.title = f"{deponent} Depo Hindsight"

    if page_margins is None:
        page_margins = {'top': 0.5, 'bottom': 0.5, 'left': 0.5, 'right': 0.5}

    section = doc.sections[0]
    section.top_margin = Inches(page_margins['top'])
    section.bottom_margin = Inches(page_margins['bottom'])
    section.left_margin = Inches(page_margins['left'])
    section.right_margin = Inches(page_margins['right'])

    GREEN = RGBColor(0, 100, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(deponent)
    run.bold = True
    run.font.size = Pt(18)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Depo Hindsight")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(100, 100, 100)

    if case_caption:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(case_caption)
        run3.italic = True
        run3.font.size = Pt(11)

    doc.add_paragraph()

    exec_sum = critique.get("executive_summary", {})
    if exec_sum:
        doc.add_heading("EXECUTIVE SUMMARY", level=1)

        findings = exec_sum.get("critical_findings", [])
        if findings:
            p = doc.add_paragraph()
            run = p.add_run("Critical Findings:")
            run.bold = True
            run.font.size = Pt(11)
            for i, finding in enumerate(findings):
                p = doc.add_paragraph(style='List Number')
                if isinstance(finding, dict):
                    text = finding.get("finding", "")
                    cite = finding.get("cite", "")
                    run = p.add_run(str(text))
                    run.font.size = Pt(11)
                    if cite:
                        run = p.add_run(f"  [{cite}]")
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 255)
                else:
                    run = p.add_run(str(finding))
                    run.font.size = Pt(11)

        cred_notes = exec_sum.get("deponent_credibility_notes", [])
        if cred_notes:
            p = doc.add_paragraph()
            run = p.add_run("Credibility Assessment:")
            run.bold = True
            run.font.size = Pt(12)
            for note in cred_notes:
                p = doc.add_paragraph()
                pattern = note.get("pattern", "")
                cite = note.get("cite", "")
                observation = note.get("observation", "")
                if pattern:
                    run = p.add_run(f"{pattern}")
                    run.bold = True
                    run.font.size = Pt(11)
                    run = p.add_run(f" --- {observation}")
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(observation)
                    run.font.size = Pt(11)
                if cite:
                    run = p.add_run(f"  [{cite}]")
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 255)

        doc.add_page_break()

    contradictions = critique.get("contradictions", [])
    if contradictions:
        doc.add_heading("CONTRADICTIONS IN TESTIMONY", level=1)
        table = doc.add_table(rows=len(contradictions) + 1, cols=5)
        table.style = 'Light Grid Accent 1'
        headers = ["Priority", "First Statement", "Contradicting Statement", "Analysis", "Trial Use"]
        widths = [Inches(0.6), Inches(1.8), Inches(1.8), Inches(1.6), Inches(1.6)]
        _add_critique_table_headers(table, headers, widths)

        for idx, item in enumerate(contradictions):
            row = table.rows[idx + 1]
            _add_priority_cell(row.cells[0], item.get("priority", ""), widths[0])

            cite1 = item.get("citation_1", "")
            text1 = item.get("testimony_1", "")
            p = row.cells[1].paragraphs[0]
            row.cells[1].width = widths[1]
            if cite1:
                run = p.add_run(f"[{cite1}] ")
                run.bold = True
                run.font.size = Pt(9)
            run = p.add_run(text1)
            run.font.size = Pt(10)

            cite2 = item.get("citation_2", "")
            text2 = item.get("testimony_2", "")
            p = row.cells[2].paragraphs[0]
            row.cells[2].width = widths[2]
            if cite2:
                run = p.add_run(f"[{cite2}] ")
                run.bold = True
                run.font.size = Pt(9)
            run = p.add_run(text2)
            run.font.size = Pt(10)

            p = row.cells[3].paragraphs[0]
            row.cells[3].width = widths[3]
            run = p.add_run(item.get("analysis", ""))
            run.font.size = Pt(10)

            p = row.cells[4].paragraphs[0]
            row.cells[4].width = widths[4]
            run = p.add_run(item.get("trial_use", ""))
            run.font.size = Pt(10)
            run.font.color.rgb = GREEN

        doc.add_page_break()

    evasions = critique.get("evasion_patterns", [])
    if evasions:
        doc.add_heading("EVASION PATTERNS", level=1)
        table = doc.add_table(rows=len(evasions) + 1, cols=5)
        table.style = 'Light Grid Accent 1'
        headers = ["Priority", "Q&A Exchange", "Evasion Type", "What Was Avoided", "Pin-Down Questions"]
        widths = [Inches(0.6), Inches(1.8), Inches(0.9), Inches(1.5), Inches(2.2)]
        _add_critique_table_headers(table, headers, widths)

        for idx, item in enumerate(evasions):
            row = table.rows[idx + 1]
            _add_priority_cell(row.cells[0], item.get("priority", ""), widths[0])

            cite = item.get("citation", "")
            q = item.get("question", "")
            a = item.get("answer", "")
            p = row.cells[1].paragraphs[0]
            row.cells[1].width = widths[1]
            if cite:
                run = p.add_run(f"[{cite}]\n")
                run.bold = True
                run.font.size = Pt(9)
            run = p.add_run(f"Q: {q}\nA: {a}")
            run.font.size = Pt(10)

            p = row.cells[2].paragraphs[0]
            row.cells[2].width = widths[2]
            run = p.add_run(item.get("evasion_type", ""))
            run.bold = True
            run.font.size = Pt(10)

            p = row.cells[3].paragraphs[0]
            row.cells[3].width = widths[3]
            run = p.add_run(item.get("what_was_avoided", ""))
            run.font.size = Pt(10)

            p = row.cells[4].paragraphs[0]
            row.cells[4].width = widths[4]
            run = p.add_run(item.get("pin_down_questions", ""))
            run.font.size = Pt(10)
            run.font.color.rgb = GREEN

        doc.add_page_break()

    commitments = critique.get("commitment_map", [])
    if commitments:
        doc.add_heading("COMMITMENT MAP", level=1)
        p = doc.add_paragraph()
        run = p.add_run("Positions the deponent locked themselves into during testimony:")
        run.italic = True
        run.font.size = Pt(10)

        table = doc.add_table(rows=len(commitments) + 1, cols=4)
        table.style = 'Light Grid Accent 1'
        headers = ["Priority", "Testimony", "Commitment", "Vulnerability"]
        widths = [Inches(0.6), Inches(2.2), Inches(2.0), Inches(2.2)]
        _add_critique_table_headers(table, headers, widths)

        for idx, item in enumerate(commitments):
            row = table.rows[idx + 1]
            _add_priority_cell(row.cells[0], item.get("priority", ""), widths[0])

            cite = item.get("citation", "")
            testimony = item.get("testimony", "")
            p = row.cells[1].paragraphs[0]
            row.cells[1].width = widths[1]
            if cite:
                run = p.add_run(f"[{cite}] ")
                run.bold = True
                run.font.size = Pt(9)
            run = p.add_run(testimony)
            run.font.size = Pt(10)

            p = row.cells[2].paragraphs[0]
            row.cells[2].width = widths[2]
            run = p.add_run(item.get("commitment", ""))
            run.bold = True
            run.font.size = Pt(10)

            p = row.cells[3].paragraphs[0]
            row.cells[3].width = widths[3]
            run = p.add_run(item.get("vulnerability", ""))
            run.font.size = Pt(10)
            run.font.color.rgb = GREEN

        doc.add_page_break()

    weak = critique.get("weak_questions", [])
    if weak:
        doc.add_heading("QUESTIONING TECHNIQUE REVIEW", level=1)
        table = doc.add_table(rows=len(weak) + 1, cols=5)
        table.style = 'Light Grid Accent 1'
        headers = ["Priority", "Citation", "Original Q&A", "Issue", "Suggested Improvement"]
        widths = [Inches(0.6), Inches(0.6), Inches(2.0), Inches(1.6), Inches(2.2)]
        _add_critique_table_headers(table, headers, widths)

        for idx, item in enumerate(weak):
            row = table.rows[idx + 1]
            _add_priority_cell(row.cells[0], item.get("priority", ""), widths[0])

            p = row.cells[1].paragraphs[0]
            row.cells[1].width = widths[1]
            run = p.add_run(item.get("citation", ""))
            run.bold = True
            run.font.size = Pt(10)

            qa_text = f"Q: {item.get('original_question', '')}\nA: {item.get('original_answer', '')}"
            p = row.cells[2].paragraphs[0]
            row.cells[2].width = widths[2]
            run = p.add_run(qa_text)
            run.font.size = Pt(10)

            p = row.cells[3].paragraphs[0]
            row.cells[3].width = widths[3]
            run = p.add_run(item.get("issue", ""))
            run.font.size = Pt(10)

            p = row.cells[4].paragraphs[0]
            row.cells[4].width = widths[4]
            run = p.add_run(item.get("suggested_improvement", ""))
            run.font.size = Pt(10)
            run.font.color.rgb = GREEN

        doc.add_page_break()

    missed = critique.get("missed_follow_ups", [])
    if missed:
        doc.add_heading("MISSED FOLLOW-UPS", level=1)
        table = doc.add_table(rows=len(missed) + 1, cols=4)
        table.style = 'Light Grid Accent 1'
        headers = ["Priority", "Testimony", "Missed Opportunity", "Suggested Follow-Up"]
        widths = [Inches(0.6), Inches(2.0), Inches(1.8), Inches(2.6)]
        _add_critique_table_headers(table, headers, widths)

        for idx, item in enumerate(missed):
            row = table.rows[idx + 1]
            _add_priority_cell(row.cells[0], item.get("priority", ""), widths[0])

            cite = item.get("citation", "")
            p = row.cells[1].paragraphs[0]
            row.cells[1].width = widths[1]
            if cite:
                run = p.add_run(f"[{cite}] ")
                run.bold = True
                run.font.size = Pt(9)
            run = p.add_run(item.get("testimony", ""))
            run.font.size = Pt(10)

            p = row.cells[2].paragraphs[0]
            row.cells[2].width = widths[2]
            run = p.add_run(item.get("missed_opportunity", ""))
            run.font.size = Pt(10)

            p = row.cells[3].paragraphs[0]
            row.cells[3].width = widths[3]
            run = p.add_run(item.get("suggested_follow_up", ""))
            run.font.size = Pt(10)
            run.font.color.rgb = GREEN

        doc.add_page_break()

    gaps = critique.get("topic_gaps", [])
    if gaps:
        doc.add_heading("TOPIC GAPS", level=1)
        p = doc.add_paragraph()
        run = p.add_run("Additional topic areas that may warrant further exploration:")
        run.italic = True
        run.font.size = Pt(10)

        for item in gaps:
            priority = item.get("priority", "")
            topic = item.get("topic", "")
            why = item.get("why_it_matters", "")
            questions = item.get("suggested_questions", [])

            p = doc.add_paragraph()
            run = p.add_run(f"[{priority}] ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = _priority_color(priority)
            run = p.add_run(topic)
            run.bold = True
            run.font.size = Pt(11)

            if why:
                p = doc.add_paragraph()
                run = p.add_run(why)
                run.font.size = Pt(10)

            if questions:
                for q in questions:
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(str(q))
                    run.font.size = Pt(10)
                    run.font.color.rgb = GREEN

            doc.add_paragraph()

        doc.add_page_break()

    ammo = critique.get("cross_examination_ammunition", [])
    if ammo:
        doc.add_heading("CROSS-EXAMINATION AMMUNITION", level=1)
        p = doc.add_paragraph()
        run = p.add_run("Testimony that can be used against the deponent at trial:")
        run.italic = True
        run.font.size = Pt(10)

        table = doc.add_table(rows=len(ammo) + 1, cols=4)
        table.style = 'Light Grid Accent 1'
        headers = ["Priority", "Testimony", "How to Use at Trial", "Setup Questions"]
        widths = [Inches(0.6), Inches(2.0), Inches(2.0), Inches(2.4)]
        _add_critique_table_headers(table, headers, widths)

        for idx, item in enumerate(ammo):
            row = table.rows[idx + 1]
            _add_priority_cell(row.cells[0], item.get("priority", ""), widths[0])

            cite = item.get("citation", "")
            p = row.cells[1].paragraphs[0]
            row.cells[1].width = widths[1]
            if cite:
                run = p.add_run(f"[{cite}] ")
                run.bold = True
                run.font.size = Pt(9)
            run = p.add_run(item.get("testimony", ""))
            run.font.size = Pt(10)

            p = row.cells[2].paragraphs[0]
            row.cells[2].width = widths[2]
            run = p.add_run(item.get("how_to_use", ""))
            run.font.size = Pt(10)

            p = row.cells[3].paragraphs[0]
            row.cells[3].width = widths[3]
            run = p.add_run(item.get("setup_questions", ""))
            run.font.size = Pt(10)
            run.font.color.rgb = GREEN

    for sec in doc.sections:
        sec.top_margin = Inches(page_margins['top'])
        sec.bottom_margin = Inches(page_margins['bottom'])
        sec.left_margin = Inches(page_margins['left'])
        sec.right_margin = Inches(page_margins['right'])

    add_header_footer(doc, logo_path)

    return doc
